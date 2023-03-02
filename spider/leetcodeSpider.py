import json
import pandas as pd
import numpy as np
import os
from docx import Document

import urllib3



# 读取题目列表数据并存入json文件中
import json
import requests
from lxml import etree

# 获取题目列表
def getQuesList(url,header):
    # fp=open("./data/questionList.json","a",encoding="utf-8")
    # 2800
    questionList=[]
    for i in range(0,2801,100):
        data={
            "query":"\n    query problemsetQuestionList($categorySlug: String, $limit: Int, $skip: Int, $filters: QuestionListFilterInput) {\n  problemsetQuestionList(\n    categorySlug: $categorySlug\n    limit: $limit\n    skip: $skip\n    filters: $filters\n  ) {\n    hasMore\n    total\n    questions {\n      acRate\n      difficulty\n      freqBar\n      frontendQuestionId\n      isFavor\n      paidOnly\n      solutionNum\n      status\n      title\n      titleCn\n      titleSlug\n      topicTags {\n        name\n        nameTranslated\n        id\n        slug\n      }\n      extra {\n        hasVideoSolution\n        topCompanyTags {\n          imgUrl\n          slug\n          numSubscribed\n        }\n      }\n    }\n  }\n}\n    ",
            "variables":{"categorySlug":"","skip":i,"limit":100,"filters":{}}
        }
        # 获取所有题目的title，存入title列表中
        response=requests.post(url=url,data=json.dumps(data),headers=header,verify=False).json()['data']["problemsetQuestionList"]['questions']
        questionList.append(response)
        # 本地存储
        # json.dump(response,fp,ensure_ascii=False)
        print(i,'条数据爬取成功！')
        # 一共2841条数据
    return questionList


# 处理quesList
# 将需要的数据存入data.csv文件中
def handleQuesList(quesList):
    title=[]
    acRate=[]
    difficulty=[]
    solutionNum=[]
    titleSlug=[]
    tagList=[]
    companyTagNum=[]
    for list in quesList:
        for item in list:
            title.append(item['frontendQuestionId']+'.'+item['titleCn'])
            acRate.append(item['acRate'])
            difficulty.append(item['difficulty'])
            solutionNum.append(item['solutionNum'])
            titleSlug.append(item['titleSlug'])
            tags=[]
            for i in item['topicTags']:
                tags.append(i['nameTranslated'])
            tagList.append(tags)
            companyTagNum.append(item['extra']['companyTagNum'])

    dic = {
        'title':title,
        'acRate':acRate,
        'difficulty':difficulty,
        'solutionNum':solutionNum,
        'titleSlug':titleSlug,
        'topicTags':tagList,
        'companyTagNum':companyTagNum
    }

    df=pd.DataFrame(dic)
    df.to_csv('../data/data.csv')

# 获取题解列表
# first：题解数量，获取100道题解
def getSolutionList(slugList,url,header):
    solutionList=[]
    for name in slugList:
        articleData={
            "operationName":"questionSolutionArticles",
            "variables":{"questionSlug":name,"first":100,"skip":0,"orderBy":"DEFAULT"},
            "query":"query questionSolutionArticles($questionSlug: String!, $skip: Int, $first: Int, $orderBy: SolutionArticleOrderBy, $userInput: String, $tagSlugs: [String!]) {\n  questionSolutionArticles(questionSlug: $questionSlug, skip: $skip, first: $first, orderBy: $orderBy, userInput: $userInput, tagSlugs: $tagSlugs) {\n    totalNum\n    edges {\n      node {\n        ...solutionArticle\n        __typename\n      }\n      __typename\n    }\n    __typename\n  }\n}\n\nfragment solutionArticle on SolutionArticleNode {\n  ipRegion\n  rewardEnabled\n  canEditReward\n  uuid\n  title\n  slug\n  sunk\n  chargeType\n  status\n  identifier\n  canEdit\n  canSee\n  reactionType\n  reactionsV2 {\n    count\n    reactionType\n    __typename\n  }\n  tags {\n    name\n    nameTranslated\n    slug\n    tagType\n    __typename\n  }\n  createdAt\n  thumbnail\n  author {\n    username\n    profile {\n      userAvatar\n      userSlug\n      realName\n      __typename\n    }\n    __typename\n  }\n  summary\n  topic {\n    id\n    commentCount\n    viewCount\n    __typename\n  }\n  byLeetcode\n  isMyFavorite\n  isMostPopular\n  isEditorsPick\n  hitCount\n  videosInfo {\n    videoId\n    coverUrl\n    duration\n    __typename\n  }\n  __typename\n}\n"
        }
        solutions=requests.post(url=url,data=json.dumps(articleData),headers=header).json()['data']['questionSolutionArticles']['edges']
        lists=[]
        for nodes in solutions:
            lists.append(nodes['node']['slug'])
        solutionList.append(lists)
    return solutionList

# 提取每个题解的content，存储为docx
def getSolutionContent(solutionList,titleList,url,header):
    for key in range(0,len(solutionList)):
        doc = Document()
        for slug in solutionList[key]:
            contentData={
                "operationName":"solutionDetailArticle",
                "variables":{"slug":slug,"orderBy":"DEFAULT"},
                "query":"query solutionDetailArticle($slug: String!, $orderBy: SolutionArticleOrderBy!) {\n  solutionArticle(slug: $slug, orderBy: $orderBy) {\n    ...solutionArticle\n    content\n    question {\n      questionTitleSlug\n      __typename\n    }\n    position\n    next {\n      slug\n      title\n      __typename\n    }\n    prev {\n      slug\n      title\n      __typename\n    }\n    __typename\n  }\n}\n\nfragment solutionArticle on SolutionArticleNode {\n  ipRegion\n  rewardEnabled\n  canEditReward\n  uuid\n  title\n  slug\n  sunk\n  chargeType\n  status\n  identifier\n  canEdit\n  canSee\n  reactionType\n  reactionsV2 {\n    count\n    reactionType\n    __typename\n  }\n  tags {\n    name\n    nameTranslated\n    slug\n    tagType\n    __typename\n  }\n  createdAt\n  thumbnail\n  author {\n    username\n    profile {\n      userAvatar\n      userSlug\n      realName\n      __typename\n    }\n    __typename\n  }\n  summary\n  topic {\n    id\n    commentCount\n    viewCount\n    __typename\n  }\n  byLeetcode\n  isMyFavorite\n  isMostPopular\n  isEditorsPick\n  hitCount\n  videosInfo {\n    videoId\n    coverUrl\n    duration\n    __typename\n  }\n  __typename\n}\n"
            }
            try:
                content=requests.post(url=url,data=json.dumps(contentData),headers=header).json()['data']['solutionArticle']['content']
                doc.add_paragraph(content)
            except:
                # 如果出现异常就跳过
                print(key,"爬取失败")
                pass
        doc.save('../data/solution/'+titleList[key]+'.docx')
        print('第',key,'条数据爬取成功！')

# 获取每道题的topicId
def getTopicId(slugList,url,header):
    topic_id=[]
    for slug in slugList:
        data={
            "operationName": "questionData",
            "variables": {"titleSlug": slug},
            "query": "query questionData($titleSlug: String!) {\n  question(titleSlug: $titleSlug) {\n    questionId\n    questionFrontendId\n    categoryTitle\n    boundTopicId\n    title\n    titleSlug\n    content\n    translatedTitle\n    translatedContent\n    isPaidOnly\n    difficulty\n    likes\n    dislikes\n    isLiked\n    similarQuestions\n    contributors {\n      username\n      profileUrl\n      avatarUrl\n      __typename\n    }\n    langToValidPlayground\n    topicTags {\n      name\n      slug\n      translatedName\n      __typename\n    }\n    companyTagStats\n    codeSnippets {\n      lang\n      langSlug\n      code\n      __typename\n    }\n    stats\n    hints\n    solution {\n      id\n      canSeeDetail\n      __typename\n    }\n    status\n    sampleTestCase\n    metaData\n    judgerAvailable\n    judgeType\n    mysqlSchemas\n    enableRunCode\n    envInfo\n    book {\n      id\n      bookName\n      pressName\n      source\n      shortDescription\n      fullDescription\n      bookImgUrl\n      pressImgUrl\n      productUrl\n      __typename\n    }\n    isSubscribed\n    isDailyQuestion\n    dailyRecordStatus\n    editorType\n    ugcQuestionId\n    style\n    exampleTestcases\n    jsonExampleTestcases\n    __typename\n  }\n}\n"
        }
        content=requests.post(url=url,data=json.dumps(data),headers=header).json()['data']['question']
        topic_id.append(content['boundTopicId'])
    return topic_id

# 获取评论
# 评论数量：first:1000 获取了1000条评论数据
def getComment(topIdList,titleList,url,header):
   for key in range(0,len(topIdList)):
        doc = Document()
        data={
            "operationName": "commonTopicComments",
            "variables": {
            "topicId": topIdList[key],
            "skip": 0,
            "orderBy": "HOT"},
            "query": "query commonTopicComments($topicId: Int!, $orderBy: CommentOrderBy, $skip: Int) {\n  commonTopicComments(topicId: $topicId, orderBy: $orderBy, skip: $skip, first:1000 ) {\n    totalNum\n    edges {\n      node {\n        ...commentFields\n        __typename\n      }\n      __typename\n    }\n    __typename\n  }\n}\n\nfragment commentFields on CommentRelayNode {\n  id\n  ipRegion\n  numChildren\n  isEdited\n  post {\n    id\n    content\n    voteUpCount\n    creationDate\n    updationDate\n    status\n    voteStatus\n    author {\n      username\n      isDiscussAdmin\n      profile {\n        userSlug\n        userAvatar\n        realName\n        __typename\n      }\n      __typename\n    }\n    mentionedUsers {\n      key\n      username\n      userSlug\n      nickName\n      __typename\n    }\n    __typename\n  }\n  __typename\n}\n"
        }
        contentList=requests.post(url=url,data=json.dumps(data),headers=header).json()['data']['commonTopicComments']['edges']
        for content in contentList:
            text=content['node']['post']['content'] 
            try:
                doc.add_paragraph(text)
            except:
                # 如果出现异常就跳过
                pass
        doc.save('../data/comment/'+titleList[key]+'.docx')
        print('第',key,'条数据爬取成功！')


def main():
    
    # 从题库中获取所有题目的数据，存入‘questionList.json’文件中
    url="https://leetcode.cn/graphql/"

    header = {
        'Origin': 'https://leetcode.cn/',
        'Accept-Encoding': 'gzip, deflate, br',
        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.109 Safari/537.36',
        'content-type': 'application/json',
        'accept': '*/*',
        'Referer': 'https://leetcode.cn/',
        'Cookie': 'did=web_62bf12a9aa82ae952a919547a686f5fa; ',
        'Connection': 'keep-alive',
        'x-original-url': 'https://leetcode.cn/'
    }

    quesList=getQuesList(url,header)
    handleQuesList(quesList)

    data = pd.read_csv('../data/data.csv')
    slugList=data['titleSlug']
    titleList=data['title']

    solutionList=getSolutionList(slugList,url,header)
    # print(len(solutionList))
    getSolutionContent(solutionList,titleList,url,header)
    topIdList=getTopicId(slugList,url,header)
    # print(len(topIdList))
    getComment(topIdList,titleList,url,header)

if __name__ == '__main__':
    urllib3.disable_warnings()
    main()
