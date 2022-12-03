import re,string

import docx
import os
import emoji

import pandas as pd
from pyltp import SentenceSplitter,Segmentor,Postagger,NamedEntityRecognizer,Parser,SementicRoleLabeller

# 处理solution
# filePath = './data/solution/'
# 处理comment
filePath = './data/comment/'
stopwordsPath = './stopwords/ltp_stopwords.txt'
LTP_DATA_DIR='D:\LTPmodel\ltp_data_v3.4.0'
cws_model_path = os.path.join(LTP_DATA_DIR, 'cws.model')  # 分词模型路径，模型名称为`cws.model`
pos_model_path = os.path.join(LTP_DATA_DIR, 'pos.model')  # 词性标注模型路径，模型名称为`pos.model`
ner_model_path = os.path.join(LTP_DATA_DIR, 'ner.model')  # 命名实体识别模型路径，模型名称为`ner.model`
par_model_path = os.path.join(LTP_DATA_DIR, 'parser.model')	 # 依存句法分析模型路径，模型名称为`parser.model`
srl_model_path = os.path.join(LTP_DATA_DIR, 'srl')	# 语义角色标注模型目录路径，模型目录为`srl`。注意该模型路径是一个目录，而不是一个文件。

# 文件操作
# 读取docx文件
def readDocx(filePath):
    fullText=[]
    # fullText=""
    doc = docx.Document(filePath)
    for p in doc.paragraphs:
        fullText.append(p.text)
        # fullText = fullText+(p.text)
    return fullText

# 读取所有文件
def findAllFile(base):
    for root,dirs,files in os.walk(base):
        for f in files:
            yield f

# 写入docx
def writeDocx(text,filePath):
    doc = docx.Document()
    p = doc.add_paragraph(text)
    doc.save(filePath)


# filePath_ls是solution文件夹下所有文件的文件名
def getAllFilePath(filePath_ls):

    for file in findAllFile(filePath):
        filePath_ls.append(file)
    # filePath_ls.sort(key=lambda x:int(x.split('.')[0]))

# 删除特殊字符
def remove_special_characters(text):
    char_pattern = re.compile(r'[\n\r\t]')

    return char_pattern.sub('',text)


# 表情符号转为字符
def emoji2code(text):
    return emoji.demojize(text)

# 删除markdown语法
def remove_markdown(text):
    markdown_pattern = re.compile(r'```.*?```|:.*?:|!\[.*?\)|[#*$\\\|]|<!.*?>|\[.*?\)')
    return markdown_pattern.sub('',text)

# 删除标点符号
def remove_punct(text):
    punct = "-：`+=(∩){},/;&.（）>!%\"'<\[\]:?；“”、！①②③④⑤？—「」^~【】@Σ∣"
    return re.sub(r"[%s]+" %punct, "",text)

# 删除字母数字下划线
def remove_num_alpha(text):
    pattern = re.compile(r'[A-Za-z0-9_]')
    return pattern.sub('',text)

# 删除空格
def remove_whiteSpace(text):
    return text.replace(" ","")


def complete_noise(data):
    new_data = remove_special_characters(data)
    new_data = emoji2code(new_data)
    new_data = remove_markdown(new_data)
    
    new_data = remove_num_alpha(new_data)
    new_data = remove_punct(new_data)
    new_data = remove_whiteSpace(new_data)
    return new_data

def create_text(list):
    result_ls = []
    for item in list:
        result_ls.append(str(complete_noise(item)))
    
    result_str = ''.join(result_ls)
    return result_str

# 分句，也就是将一片文本分割为独立的句子
def sentence_splitter(sentence):
	sents = SentenceSplitter.split(sentence)  # 分句
	return sents

def segmentor(sentence):
	segmentor = Segmentor(cws_model_path)	 # 初始化实例
	# segmentor.load(cws_model_path)	# 加载模型
	#segmentor.load_with_lexicon('cws_model_path', 'D:\pyprojects\LTP\ltp_data\dict.txt') #加载模型	  使用用户自定义字典的高级分词
	words = segmentor.segment(sentence)	 # 分词
	segmentor.release()	 # 释放模型
	return words


#创建停用词表
def remove_stopwords_punct(text,path):
	result_ls = []
	stopwords = [line.strip() for line in open(path, 'r', encoding='utf-8').readlines()]
	for word in text:
		if word not in stopwords:
			result_ls.append(str(word))
	result_str = ''.join(result_ls)
	return result_str

def main():
    filePath_ls=[]
    getAllFilePath(filePath_ls)
    for item in filePath_ls:
        path = filePath+item
        list=readDocx(path)
        doc = docx.Document()
        data=create_text(list)
        sent = sentence_splitter(data)
        for s in sent:
            seg = segmentor(s)
            seg = remove_stopwords_punct(seg,stopwordsPath)
            p = doc.add_paragraph(seg)
        doc.save(path)
        print(item,'write successfully!')

if __name__ == '__main__':
    main()